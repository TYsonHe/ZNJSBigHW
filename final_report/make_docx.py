from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont

import experiment

BASE_DIR = Path(__file__).resolve().parent
random_instance = experiment.random_instance


def set_east_asian_font(run, font_name: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for name, size, font in [("Heading 1", 15, "黑体"), ("Heading 2", 13, "黑体")]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_page_number(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = text.replace("`", "")
    text = text.replace("$", "")
    text = text.replace("\\ldots", "…")
    text = text.replace("\\sum", "∑")
    text = text.replace("\\quad", "，")
    text = text.replace("\\alpha", "α")
    text = text.replace("\\beta", "β")
    text = text.replace("\\tau", "τ")
    text = text.replace("\\eta", "η")
    return text.strip()


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(clean_inline(text))
    set_east_asian_font(run)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(clean_inline(text), level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_east_asian_font(run, "黑体")


def add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run(clean_inline(text))
    run.bold = True
    run.font.size = Pt(18)
    set_east_asian_font(run, "黑体")


def add_summary_table(doc: Document, summary_csv: Path) -> None:
    with summary_csv.open(encoding="utf-8") as f:
        rows = list(csv.reader(f))
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_east_asian_font(run)
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run("表 1  ACO 与 GA 实验结果汇总")
    run.font.size = Pt(10)
    set_east_asian_font(run, "宋体")


def load_font(size: int):
    for name in ["simhei.ttf", "msyh.ttc", "simsun.ttc"]:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def render_summary_chart(summary_csv: Path, png_path: Path) -> None:
    with summary_csv.open(encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    values = [(f"{row['instance']}\n{row['algorithm']}", float(row["mean_length"])) for row in rows]
    width, height = 1200, 720
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34)
    label_font = load_font(20)
    axis_font = load_font(18)
    draw.text((width // 2, 35), "ACO 与 GA 平均路径长度对比", fill=(0, 0, 0), font=title_font, anchor="mm")
    left, top, right, bottom = 120, 110, 1120, 590
    draw.line((left, bottom, right, bottom), fill=(0, 0, 0), width=2)
    draw.line((left, top, left, bottom), fill=(0, 0, 0), width=2)
    max_value = max(value for _, value in values) * 1.1
    colors = {"ACO": (31, 119, 180), "GA": (214, 39, 40)}
    bar_gap = 50
    bar_width = (right - left - bar_gap * (len(values) + 1)) / len(values)
    for index, (label, value) in enumerate(values):
        x0 = left + bar_gap + index * (bar_width + bar_gap)
        x1 = x0 + bar_width
        y0 = bottom - value / max_value * (bottom - top)
        color = colors.get(label.split("\n")[-1], (80, 80, 80))
        draw.rectangle((x0, y0, x1, bottom), fill=color)
        draw.text(((x0 + x1) / 2, y0 - 16), f"{value:.1f}", fill=(0, 0, 0), font=axis_font, anchor="mm")
        draw.text(((x0 + x1) / 2, bottom + 34), label, fill=(0, 0, 0), font=label_font, anchor="mm", align="center")
    draw.text((60, (top + bottom) / 2), "平均路径长度", fill=(0, 0, 0), font=axis_font, anchor="mm")
    draw.text((width // 2, 670), "数据来源：final_report/results/summary.csv", fill=(80, 80, 80), font=axis_font, anchor="mm")
    image.save(png_path)


def run_algorithm_for_chart(name: str, points: list[tuple[float, float]], seed: int, iterations: int):
    if name == "ACO":
        best, route, curve = experiment.run_aco(points, seed, iterations=iterations, ants=30)
    else:
        best, route, curve = experiment.run_ga(points, seed, generations=iterations, pop_size=50)
    return {"name": name, "best": best, "route": route, "curve": curve}


def render_convergence_chart(series: list[dict], png_path: Path) -> None:
    width, height = 1200, 720
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34)
    label_font = load_font(20)
    draw.text((width // 2, 35), "Random-30 实例收敛曲线", fill=(0, 0, 0), font=title_font, anchor="mm")
    left, top, right, bottom = 100, 100, 1120, 600
    draw.line((left, bottom, right, bottom), fill=(0, 0, 0), width=2)
    draw.line((left, top, left, bottom), fill=(0, 0, 0), width=2)
    values = [value for item in series for value in item["curve"]]
    min_value, max_value = min(values), max(values)
    colors = {"ACO": (31, 119, 180), "GA": (214, 39, 40)}
    for item in series:
        points = []
        curve = item["curve"]
        for index, value in enumerate(curve):
            x = left + index * (right - left) / max(1, len(curve) - 1)
            y = bottom - (value - min_value) * (bottom - top) / (max_value - min_value + 1e-9)
            points.append((x, y))
        draw.line(points, fill=colors[item["name"]], width=4)
        draw.text((right - 120, top + 35 * len([s for s in series if s["name"] <= item["name"]])), item["name"], fill=colors[item["name"]], font=label_font)
    draw.text((width // 2, 665), "横轴：迭代次数；纵轴：当前最优路径长度", fill=(80, 80, 80), font=label_font, anchor="mm")
    image.save(png_path)


def scale_for_chart(points: list[tuple[float, float]], box: tuple[int, int, int, int]) -> list[tuple[float, float]]:
    left, top, right, bottom = box
    xs = [p[0] for p in points]
    ys = [p[1] for p in points]
    min_x, max_x = min(xs), max(xs)
    min_y, max_y = min(ys), max(ys)
    return [
        (
            left + (x - min_x) * (right - left) / (max_x - min_x + 1e-9),
            bottom - (y - min_y) * (bottom - top) / (max_y - min_y + 1e-9),
        )
        for x, y in points
    ]


def render_route_chart(points: list[tuple[float, float]], series: list[dict], png_path: Path) -> None:
    width, height = 1200, 720
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34)
    label_font = load_font(20)
    draw.text((width // 2, 35), "Random-30 实例最优路径对比", fill=(0, 0, 0), font=title_font, anchor="mm")
    boxes = [(70, 110, 560, 610), (640, 110, 1130, 610)]
    colors = {"ACO": (31, 119, 180), "GA": (214, 39, 40)}
    for box, item in zip(boxes, series):
        scaled = scale_for_chart(points, box)
        route_points = [scaled[index] for index in item["route"]] + [scaled[item["route"][0]]]
        draw.rectangle(box, outline=(120, 120, 120), width=2)
        draw.line(route_points, fill=colors[item["name"]], width=3)
        for x, y in scaled:
            draw.ellipse((x - 4, y - 4, x + 4, y + 4), fill=(0, 0, 0))
        draw.text(((box[0] + box[2]) / 2, 640), f"{item['name']}：最优长度 {item['best']:.2f}", fill=colors[item["name"]], font=label_font, anchor="mm")
    image.save(png_path)


def add_image(doc: Document, image_path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(14.5))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    caption_run.font.size = Pt(10)
    set_east_asian_font(caption_run)


def markdown_to_docx(doc: Document, markdown: str, summary_csv: Path, images: list[Path]) -> None:
    in_table = False
    skip_formula = False
    image_index = 0
    inserted_table = False
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line == "$$":
            skip_formula = not skip_formula
            continue
        if skip_formula:
            add_paragraph(doc, "路径总长度 L(R)=∑ d(r_k,r_{k+1})，其中 r_{n+1}=r_1。")
            skip_formula = False
            continue
        if line.startswith("|"):
            in_table = True
            continue
        if in_table:
            in_table = False
        if line.startswith("# "):
            add_title(doc, line[2:])
        elif line.startswith("## "):
            heading = line[3:]
            add_heading(doc, heading, 1)
            if heading.startswith("5 实验结果") and not inserted_table:
                add_paragraph(doc, "表 1 汇总了两种算法在两个实例上的最优路径长度、平均路径长度、标准差和平均运行时间。")
                add_summary_table(doc, summary_csv)
                inserted_table = True
        elif line.startswith("### "):
            add_heading(doc, line[4:], 2)
        elif line.startswith("- "):
            add_paragraph(doc, "（1）" + line[2:])
        elif line.startswith("[1]") or line.startswith("[2]") or line.startswith("[3]") or line.startswith("[4]"):
            add_paragraph(doc, line)
        elif "results/convergence.svg" in line and image_index < len(images):
            add_paragraph(doc, clean_inline(line))
            add_image(doc, images[image_index], "图 1  Random-30 实例下 ACO 与 GA 的平均收敛曲线")
            image_index += 1
        elif "results/routes.svg" in line and image_index < len(images):
            add_paragraph(doc, clean_inline(line))
            add_image(doc, images[image_index], "图 2  Random-30 实例下 ACO 与 GA 的最优路径对比")
            image_index += 1
        else:
            add_paragraph(doc, line)

    while image_index < len(images):
        add_image(doc, images[image_index], f"图 {image_index + 1}  实验图表")
        image_index += 1


def create_docx(source: Path, output: Path) -> None:
    summary = BASE_DIR / "results" / "summary.csv"
    if not summary.exists():
        summary = source.parent / "summary.csv"
    create_polished_docx(source, summary, output, [])


def create_polished_docx(source: Path, summary_csv: Path, output: Path, images: list[Path]) -> None:
    doc = Document()
    configure_document(doc)
    add_page_number(doc.sections[0])
    markdown_to_docx(doc, source.read_text(encoding="utf-8"), summary_csv, images)
    doc.save(output)


def main() -> None:
    summary_png = BASE_DIR / "results" / "summary_chart.png"
    convergence_png = BASE_DIR / "results" / "convergence_chart.png"
    routes_png = BASE_DIR / "results" / "routes_chart.png"
    points = experiment.random_instance(30, 20260608)
    chart_series = [
        run_algorithm_for_chart("ACO", points, 0, 120),
        run_algorithm_for_chart("GA", points, 0, 120),
    ]
    render_summary_chart(BASE_DIR / "results" / "summary.csv", summary_png)
    render_convergence_chart(chart_series, convergence_png)
    render_route_chart(points, chart_series, routes_png)
    create_polished_docx(
        BASE_DIR / "report.md",
        BASE_DIR / "results" / "summary.csv",
        BASE_DIR / "面向物流配送路径规划的蚁群算法与遗传算法比较研究.docx",
        [summary_png, convergence_png, routes_png],
    )
    print("wrote polished docx")


if __name__ == "__main__":
    main()
