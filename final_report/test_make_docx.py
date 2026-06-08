import csv
import tempfile
import unittest
from pathlib import Path

from docx import Document
from PIL import Image

import make_docx


class MakeDocxTest(unittest.TestCase):
    def test_create_polished_docx_with_table_and_images(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            source = temp / "report.md"
            summary = temp / "summary.csv"
            image = temp / "chart.png"
            output = temp / "report.docx"
            source.write_text(
                "# 测试标题\n\n## 摘要\n\n这是润色后的正文。\n\n## 5 实验结果与分析\n\n这里应插入实验表格。",
                encoding="utf-8",
            )
            with summary.open("w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow(["instance", "algorithm", "best_length", "mean_length", "std_length", "mean_time_seconds"])
                writer.writerow(["random_30", "ACO", "441.1274", "442.7554", "1.0976", "1.2682"])
                writer.writerow(["random_30", "GA", "488.6986", "545.7955", "34.1085", "0.4380"])
            Image.new("RGB", (320, 180), "white").save(image)

            make_docx.create_polished_docx(source, summary, output, [image])

            self.assertTrue(output.exists())
            doc = Document(output)
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("测试标题", text)
            self.assertIn("这是润色后的正文。", text)
            self.assertIn("图 1", text)
            self.assertGreaterEqual(len(doc.tables), 1)
            self.assertGreaterEqual(len(doc.inline_shapes), 1)

    def test_render_summary_chart_uses_data(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            summary = temp / "summary.csv"
            output = temp / "summary_chart.png"
            with summary.open("w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow(["instance", "algorithm", "best_length", "mean_length", "std_length", "mean_time_seconds"])
                writer.writerow(["random_30", "ACO", "441.1274", "442.7554", "1.0976", "1.2682"])
                writer.writerow(["random_30", "GA", "488.6986", "545.7955", "34.1085", "0.4380"])
            make_docx.render_summary_chart(summary, output)
            self.assertTrue(output.exists())
            image = Image.open(output)
            colors = image.getcolors(maxcolors=1_000_000)
            self.assertIsNotNone(colors)
            self.assertGreater(len(colors), 10)

    def test_render_curve_and_route_images_use_experiment_data(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            curve = temp / "curve.png"
            route = temp / "route.png"
            points = make_docx.random_instance(12, 1)
            aco = make_docx.run_algorithm_for_chart("ACO", points, 1, 8)
            ga = make_docx.run_algorithm_for_chart("GA", points, 1, 8)
            make_docx.render_convergence_chart([aco, ga], curve)
            make_docx.render_route_chart(points, [aco, ga], route)
            self.assertTrue(curve.exists())
            self.assertTrue(route.exists())
            self.assertGreater(len(Image.open(curve).getcolors(maxcolors=1_000_000)), 10)
            self.assertGreater(len(Image.open(route).getcolors(maxcolors=1_000_000)), 10)


if __name__ == "__main__":
    unittest.main()
